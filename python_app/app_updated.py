# -*- coding: utf-8 -*-
"""
DOCX Formatting Validator & Comparator - Python Local Server Code

A Flask-based application for scanning, validating, and comparing the formatting of .docx documents.
Supports file upload OR file path input.

Features:
  - File upload support (Choose File button)
  - Font color detection
  - Paragraph alignment check
  - Indentation validation
  - Bullet / numbering detection (XML-based)
  - Bullet level consistency (alignment hierarchy)
  - Table comparison
  - Spacing validation
  - 3-file comparison: WRD vs Handoff + WRD vs Support

Requires:
    pip install flask python-docx lxml

Usage:
    python app_updated.py
    Then direct your browser to http://127.0.0.1:5000
"""

import os
import json
import csv
import re
import tempfile
import shutil
from pathlib import Path
from datetime import datetime
from flask import Flask, request, jsonify, render_template_string

# Try to import python-docx
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Warning: 'python-docx' library is not installed. Install with 'pip install python-docx'")

app = Flask(__name__)

# =============================================================================
# PREDEFINED STANDARD RULES
# =============================================================================
DEFAULT_STANDARDS = [
    {
        "id": "heading1",
        "name": "Level 1 Heading",
        "fontName": "Human Black Condensed",
        "fontSize": 20.0,
        "bold": True,
        "alignment": "LEFT",
        "indent": 0.0
    },
    {
        "id": "heading2",
        "name": "Level 2 Heading",
        "fontName": "Human Black Condensed",
        "fontSize": 11.5,
        "bold": True,
        "alignment": "LEFT",
        "indent": 0.0
    },
    {
        "id": "heading3",
        "name": "Level 3 Heading",
        "fontName": "Human Black Condensed",
        "fontSize": 10.5,
        "bold": True,
        "alignment": "LEFT",
        "indent": 0.0
    },
    {
        "id": "normal",
        "name": "Normal Text",
        "fontName": "Human Light Condensed",
        "fontSize": 10.5,
        "bold": False,
        "alignment": "LEFT",
        "indent": 0.0
    },
    {
        "id": "boldText",
        "name": "Text Level Bold",
        "fontName": "Human Bold Condensed",
        "fontSize": 10.5,
        "bold": True,
        "alignment": "LEFT",
        "indent": 0.0
    },
    {
        "id": "formName",
        "name": "Form Name",
        "fontName": "Human Light Condensed",
        "fontSize": 7.0,
        "bold": False,
        "alignment": "LEFT",
        "indent": 0.0
    }
]

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# =============================================================================
# Helper functions for WRD comparison
# =============================================================================

def get_alignment_name(align):
    """Convert WD_ALIGN_PARAGRAPH enum to readable string."""
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
        WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
        WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
    }
    return mapping.get(align, "UNKNOWN")


def get_font_color(run):
    """Extract font color as hex string (e.g. 'FF0000' for red)."""
    try:
        if run.font.color and run.font.color.rgb:
            return str(run.font.color.rgb)
    except Exception:
        pass
    # Also check XML-level color
    try:
        rpr = run._element.find(f"{W_NS}rPr")
        if rpr is not None:
            color_elem = rpr.find(f"{W_NS}color")
            if color_elem is not None:
                val = color_elem.get(f"{W_NS}val")
                if val:
                    return val
    except Exception:
        pass
    return None


def get_highlight_color(run):
    """Extract highlight/background color from a run."""
    try:
        rpr = run._element.find(f"{W_NS}rPr")
        if rpr is not None:
            highlight = rpr.find(f"{W_NS}highlight")
            if highlight is not None:
                val = highlight.get(f"{W_NS}val")
                if val:
                    return val
    except Exception:
        pass
    return None


# ===== FIXED: Using findall instead of xpath to avoid 'namespaces' error =====

def is_bullet(para):
    """Check if paragraph is a bullet or numbered list item."""
    num_pr = para._element.findall(f".//{W_NS}numPr")
    return len(num_pr) > 0


def get_bullet_level(para):
    """Get the bullet/numbering indent level (0=main, 1=sub, 2=sub-sub)."""
    ilvl_elems = para._element.findall(f".//{W_NS}ilvl")
    if ilvl_elems:
        elem = ilvl_elems[0]
        # Try with namespace prefix first, then without
        val = elem.get(f"{W_NS}val")
        if val is None:
            val = elem.get("val")
        if val is not None:
            try:
                return int(val)
            except (ValueError, TypeError):
                return 0
    return None


def get_num_id(para):
    """Get the numbering definition ID (to distinguish bullet vs numbered list)."""
    num_id_elems = para._element.findall(f".//{W_NS}numId")
    if num_id_elems:
        elem = num_id_elems[0]
        val = elem.get(f"{W_NS}val")
        if val is None:
            val = elem.get("val")
        return val
    return None


def get_underline(run):
    """Check if run has underline formatting."""
    try:
        return run.underline
    except Exception:
        return None


def get_italic(run):
    """Check if run has italic formatting."""
    try:
        return run.italic
    except Exception:
        return None


def get_strikethrough(run):
    """Check if run has strikethrough formatting."""
    try:
        rpr = run._element.find(f"{W_NS}rPr")
        if rpr is not None:
            strike = rpr.find(f"{W_NS}strike")
            if strike is not None:
                val = strike.get(f"{W_NS}val")
                return val is None or val.lower() in ("true", "1")
        return False
    except Exception:
        return False


# =============================================================================
# Document defaults parser
# =============================================================================

def get_doc_defaults(doc):
    """Parse styles.xml to get default document properties."""
    defaults = {"fontName": None, "fontSize": None, "bold": None}
    try:
        styles_element = doc.styles.element
        doc_def = styles_element.find(f"{W_NS}docDefaults")
        if doc_def is not None:
            rpr_default = doc_def.find(f"{W_NS}rPrDefault")
            if rpr_default is not None:
                rpr = rpr_default.find(f"{W_NS}rPr")
                if rpr is not None:
                    rfonts = rpr.find(f"{W_NS}rFonts")
                    if rfonts is not None:
                        for attr in ["ascii", "hAnsi", "cs"]:
                            val = rfonts.get(f"{W_NS}{attr}")
                            if val:
                                defaults["fontName"] = val
                                break
                    sz = rpr.find(f"{W_NS}sz")
                    if sz is not None:
                        hp = sz.get(f"{W_NS}val")
                        if hp:
                            defaults["fontSize"] = int(hp) / 2.0
                    b_elem = rpr.find(f"{W_NS}b")
                    if b_elem is not None:
                        b_val = b_elem.get(f"{W_NS}val")
                        defaults["bold"] = b_val is None or b_val.lower() in ("true", "1")
    except Exception:
        pass
    return defaults


def get_xml_run_props(r_elem):
    """Extract formatting values from a single run element."""
    props = {"fontName": None, "fontSize": None, "bold": None, "runStyleId": None}
    rpr = r_elem.find(f"{W_NS}rPr")
    if rpr is None:
        return props

    rfonts = rpr.find(f"{W_NS}rFonts")
    if rfonts is not None:
        for attr in ["ascii", "hAnsi", "cs"]:
            val = rfonts.get(f"{W_NS}{attr}")
            if val:
                props["fontName"] = val
                break

    sz = rpr.find(f"{W_NS}sz")
    if sz is not None:
        hp = sz.get(f"{W_NS}val")
        if hp:
            props["fontSize"] = int(hp) / 2.0

    b_elem = rpr.find(f"{W_NS}b")
    if b_elem is not None:
        b_val = b_elem.get(f"{W_NS}val")
        props["bold"] = b_val is None or b_val.lower() in ("true", "1")

    rstyle = rpr.find(f"{W_NS}rStyle")
    if rstyle is not None:
        props["runStyleId"] = rstyle.get(f"{W_NS}val")

    return props


def walk_style_chain(style, key_func, max_depth=10):
    """Traverse style hierarchy (based_on) to resolve formatting."""
    current = style
    depth = 0
    while current and depth < max_depth:
        val = key_func(current)
        if val is not None:
            return val
        try:
            current = current.base_style
        except Exception:
            break
        depth += 1
    return None


def resolve_properties(run_props, para, doc_defaults, doc_styles):
    """Resolve final font properties based on XML cascade."""
    font_name = run_props["fontName"]
    font_size = run_props["fontSize"]
    bold = run_props["bold"]

    para_style = para.style
    if para_style:
        if font_name is None:
            font_name = walk_style_chain(
                para_style,
                lambda s: s.font.name if s.font and s.font.name else None
            )
        if font_size is None:
            font_size = walk_style_chain(
                para_style,
                lambda s: s.font.size.pt if s.font and s.font.size else None
            )
        if bold is None:
            bold = walk_style_chain(
                para_style,
                lambda s: s.font.bold if s.font and s.font.bold is not None else None
            )

    if font_name is None:
        font_name = doc_defaults.get("fontName")
    if font_size is None:
        font_size = doc_defaults.get("fontSize")
    if bold is None:
        bold = doc_defaults.get("bold", False)

    return {"fontName": font_name, "fontSize": font_size, "bold": bold}


# =============================================================================
# Paragraph analysis
# =============================================================================

def analyze_paragraph(para, doc_defaults=None, doc_styles=None):
    """Extract ALL formatting details from a paragraph."""
    pf = para.paragraph_format

    para_data = {
        "text": para.text.strip(),
        "style_name": para.style.name if para.style else None,
        "alignment": get_alignment_name(para.alignment),
        "indent_left": pf.left_indent.pt if pf.left_indent else 0,
        "indent_right": pf.right_indent.pt if pf.right_indent else 0,
        "first_line_indent": pf.first_line_indent.pt if pf.first_line_indent else 0,
        "is_bullet": is_bullet(para),
        "bullet_level": get_bullet_level(para),
        "num_id": get_num_id(para),
        "space_before": pf.space_before.pt if pf.space_before else None,
        "space_after": pf.space_after.pt if pf.space_after else None,
        "line_spacing": pf.line_spacing if pf.line_spacing else None,
        "runs": []
    }

    for run in para.runs:
        xml_props = get_xml_run_props(run._element)
        resolved = resolve_properties(xml_props, para, doc_defaults or {}, doc_styles or {})

        run_data = {
            "text": run.text,
            "font_name": run.font.name or resolved.get("fontName"),
            "font_size": run.font.size.pt if run.font.size else resolved.get("fontSize"),
            "bold": run.bold if run.bold is not None else resolved.get("bold", False),
            "italic": get_italic(run),
            "underline": get_underline(run),
            "strikethrough": get_strikethrough(run),
            "color": get_font_color(run),
            "highlight": get_highlight_color(run),
        }
        para_data["runs"].append(run_data)

    return para_data


# =============================================================================
# Table analysis
# =============================================================================

def analyze_table(table, doc_defaults=None, doc_styles=None):
    """Extract ALL data and formatting from a table."""
    table_data = {
        "rows": len(table.rows),
        "cols": len(table.columns),
        "cells": []
    }

    for r_idx, row in enumerate(table.rows):
        row_data = []
        for c_idx, cell in enumerate(row.cells):
            cell_data = {
                "row": r_idx,
                "col": c_idx,
                "text": cell.text.strip(),
                "paragraphs": []
            }
            for para in cell.paragraphs:
                p_data = analyze_paragraph(para, doc_defaults, doc_styles)
                cell_data["paragraphs"].append(p_data)

            row_data.append(cell_data)
        table_data["cells"].append(row_data)

    return table_data


# =============================================================================
# Complete document analysis
# =============================================================================

def analyze_docx_file(filepath, standards=None):
    """Extract ALL paragraphs + tables with FULL formatting details."""
    if standards is None:
        standards = DEFAULT_STANDARDS

    doc = Document(str(filepath))
    doc_defaults = get_doc_defaults(doc)
    doc_styles = {}

    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        para_data = analyze_paragraph(para, doc_defaults, doc_styles)
        para_data["para_id"] = i
        paragraphs.append(para_data)

    tables = []
    for t_idx, table in enumerate(doc.tables):
        t_data = analyze_table(table, doc_defaults, doc_styles)
        t_data["table_id"] = t_idx
        tables.append(t_data)

    validation_issues = validate_against_standards(paragraphs, standards)

    return {
        "file": str(filepath),
        "total_paragraphs": len(paragraphs),
        "total_tables": len(tables),
        "paragraphs": paragraphs,
        "tables": tables,
        "validation_issues": validation_issues,
        "analyzed_at": datetime.now().isoformat()
    }


# =============================================================================
# Validate paragraphs against formatting standards
# =============================================================================

def validate_against_standards(paragraphs, standards):
    """Check each paragraph against predefined formatting rules."""
    issues = []

    for para in paragraphs:
        para_issues = []

        matched_standard = None
        for std in standards:
            if std["name"].lower() in (para.get("style_name") or "").lower():
                matched_standard = std
                break

        if matched_standard is None:
            continue

        if para["alignment"] != matched_standard.get("alignment", "LEFT"):
            para_issues.append({
                "type": "ALIGNMENT_MISMATCH",
                "expected": matched_standard.get("alignment"),
                "found": para["alignment"]
            })

        if abs(para["indent_left"] - matched_standard.get("indent", 0)) > 0.5:
            para_issues.append({
                "type": "INDENT_MISMATCH",
                "expected": matched_standard.get("indent"),
                "found": para["indent_left"]
            })

        for run in para["runs"]:
            if run["font_name"] and run["font_name"] != matched_standard.get("fontName"):
                para_issues.append({
                    "type": "FONT_MISMATCH",
                    "expected": matched_standard.get("fontName"),
                    "found": run["font_name"]
                })
            if run["font_size"] and abs(run["font_size"] - matched_standard.get("fontSize", 0)) > 0.1:
                para_issues.append({
                    "type": "FONT_SIZE_MISMATCH",
                    "expected": matched_standard.get("fontSize"),
                    "found": run["font_size"]
                })
            if run["bold"] != matched_standard.get("bold", False):
                para_issues.append({
                    "type": "BOLD_MISMATCH",
                    "expected": matched_standard.get("bold"),
                    "found": run["bold"]
                })

        if para_issues:
            issues.append({
                "para_id": para.get("para_id"),
                "text_preview": para["text"][:80],
                "issues": para_issues
            })

    return issues


# =============================================================================
# Document comparison engine (two documents)
# =============================================================================

def compare_documents(doc1_data, doc2_data):
    """Compare two analyzed documents paragraph-by-paragraph and table-by-table."""
    results = {
        "paragraph_issues": [],
        "table_issues": [],
        "summary": {}
    }

    p1_list = doc1_data["paragraphs"]
    p2_list = doc2_data["paragraphs"]
    max_para = max(len(p1_list), len(p2_list))

    total_issues = 0

    for i in range(max_para):
        row = {"para_id": i, "issues": []}

        if i >= len(p1_list):
            row["issues"].append({"type": "MISSING_IN_WRD", "detail": f"Paragraph {i} exists in Reference but missing in WRD"})
            results["paragraph_issues"].append(row)
            total_issues += 1
            continue

        if i >= len(p2_list):
            row["issues"].append({"type": "MISSING_IN_REFERENCE", "detail": f"Paragraph {i} exists in WRD but missing in Reference"})
            results["paragraph_issues"].append(row)
            total_issues += 1
            continue

        p1 = p1_list[i]
        p2 = p2_list[i]

        if p1["text"] != p2["text"]:
            row["issues"].append({"type": "TEXT_MISMATCH", "wrd": p1["text"][:100], "reference": p2["text"][:100]})

        if p1["alignment"] != p2["alignment"]:
            row["issues"].append({"type": "ALIGNMENT_MISMATCH", "wrd": p1["alignment"], "reference": p2["alignment"]})

        if p1["indent_left"] != p2["indent_left"]:
            row["issues"].append({"type": "INDENT_LEFT_MISMATCH", "wrd": p1["indent_left"], "reference": p2["indent_left"]})

        if p1["first_line_indent"] != p2["first_line_indent"]:
            row["issues"].append({"type": "FIRST_LINE_INDENT_MISMATCH", "wrd": p1["first_line_indent"], "reference": p2["first_line_indent"]})

        if p1["is_bullet"] != p2["is_bullet"]:
            row["issues"].append({"type": "BULLET_PRESENCE_MISMATCH", "wrd": p1["is_bullet"], "reference": p2["is_bullet"]})

        if p1["bullet_level"] != p2["bullet_level"]:
            row["issues"].append({"type": "BULLET_LEVEL_MISMATCH", "wrd": p1["bullet_level"], "reference": p2["bullet_level"]})

        if p1["space_before"] != p2["space_before"]:
            row["issues"].append({"type": "SPACE_BEFORE_MISMATCH", "wrd": p1["space_before"], "reference": p2["space_before"]})

        if p1["space_after"] != p2["space_after"]:
            row["issues"].append({"type": "SPACE_AFTER_MISMATCH", "wrd": p1["space_after"], "reference": p2["space_after"]})

        max_runs = max(len(p1["runs"]), len(p2["runs"]))
        for r_idx in range(max_runs):
            if r_idx >= len(p1["runs"]):
                row["issues"].append({"type": "EXTRA_RUN_IN_REFERENCE", "run_index": r_idx})
                continue
            if r_idx >= len(p2["runs"]):
                row["issues"].append({"type": "EXTRA_RUN_IN_WRD", "run_index": r_idx})
                continue

            r1 = p1["runs"][r_idx]
            r2 = p2["runs"][r_idx]

            if r1["font_name"] != r2["font_name"]:
                row["issues"].append({"type": "FONT_NAME_MISMATCH", "run_index": r_idx, "wrd": r1["font_name"], "reference": r2["font_name"]})
            if r1["font_size"] != r2["font_size"]:
                row["issues"].append({"type": "FONT_SIZE_MISMATCH", "run_index": r_idx, "wrd": r1["font_size"], "reference": r2["font_size"]})
            if r1["bold"] != r2["bold"]:
                row["issues"].append({"type": "BOLD_MISMATCH", "run_index": r_idx, "wrd": r1["bold"], "reference": r2["bold"]})
            if r1["italic"] != r2["italic"]:
                row["issues"].append({"type": "ITALIC_MISMATCH", "run_index": r_idx, "wrd": r1["italic"], "reference": r2["italic"]})
            if r1["underline"] != r2["underline"]:
                row["issues"].append({"type": "UNDERLINE_MISMATCH", "run_index": r_idx, "wrd": r1["underline"], "reference": r2["underline"]})
            if r1["strikethrough"] != r2["strikethrough"]:
                row["issues"].append({"type": "STRIKETHROUGH_MISMATCH", "run_index": r_idx, "wrd": r1["strikethrough"], "reference": r2["strikethrough"]})
            if r1["color"] != r2["color"]:
                row["issues"].append({"type": "FONT_COLOR_MISMATCH", "run_index": r_idx, "wrd": r1["color"], "reference": r2["color"]})
            if r1["highlight"] != r2["highlight"]:
                row["issues"].append({"type": "HIGHLIGHT_COLOR_MISMATCH", "run_index": r_idx, "wrd": r1["highlight"], "reference": r2["highlight"]})

        if row["issues"]:
            total_issues += len(row["issues"])
            results["paragraph_issues"].append(row)

    # Compare tables
    t1_list = doc1_data["tables"]
    t2_list = doc2_data["tables"]
    max_tables = max(len(t1_list), len(t2_list))

    for t_idx in range(max_tables):
        t_row = {"table_id": t_idx, "issues": []}

        if t_idx >= len(t1_list):
            t_row["issues"].append({"type": "TABLE_MISSING_IN_WRD", "detail": f"Table {t_idx} in Reference but missing in WRD"})
            results["table_issues"].append(t_row)
            total_issues += 1
            continue
        if t_idx >= len(t2_list):
            t_row["issues"].append({"type": "TABLE_MISSING_IN_REFERENCE", "detail": f"Table {t_idx} in WRD but missing in Reference"})
            results["table_issues"].append(t_row)
            total_issues += 1
            continue

        t1 = t1_list[t_idx]
        t2 = t2_list[t_idx]

        if t1["rows"] != t2["rows"]:
            t_row["issues"].append({"type": "TABLE_ROW_COUNT_MISMATCH", "wrd": t1["rows"], "reference": t2["rows"]})
        if t1["cols"] != t2["cols"]:
            t_row["issues"].append({"type": "TABLE_COL_COUNT_MISMATCH", "wrd": t1["cols"], "reference": t2["cols"]})

        min_rows = min(len(t1["cells"]), len(t2["cells"]))
        for r_idx in range(min_rows):
            min_cols = min(len(t1["cells"][r_idx]), len(t2["cells"][r_idx]))
            for c_idx in range(min_cols):
                c1 = t1["cells"][r_idx][c_idx]
                c2 = t2["cells"][r_idx][c_idx]
                if c1["text"] != c2["text"]:
                    t_row["issues"].append({"type": "TABLE_CELL_TEXT_MISMATCH", "row": r_idx, "col": c_idx, "wrd": c1["text"][:80], "reference": c2["text"][:80]})

        if t_row["issues"]:
            total_issues += len(t_row["issues"])
            results["table_issues"].append(t_row)

    results["summary"] = {
        "wrd_paragraphs": len(p1_list),
        "ref_paragraphs": len(p2_list),
        "wrd_tables": len(t1_list),
        "ref_tables": len(t2_list),
        "total_issues_found": total_issues,
        "paragraphs_with_issues": len(results["paragraph_issues"]),
        "tables_with_issues": len(results["table_issues"]),
        "compared_at": datetime.now().isoformat()
    }

    return results


# =============================================================================
# 3-file comparison (WRD vs Handoff + WRD vs Support)
# =============================================================================

def compare_three_documents(wrd_data, handoff_data=None, support_data=None):
    """Compare WRD against Handoff and/or Support reference documents."""
    result = {
        "wrd_file": wrd_data.get("file", ""),
        "comparisons": [],
        "combined_summary": {}
    }

    total_all_issues = 0
    handoff_issues = 0
    support_issues = 0

    if handoff_data:
        handoff_result = compare_documents(wrd_data, handoff_data)
        handoff_result["comparison_label"] = "WRD vs Handoff"
        handoff_result["reference_file"] = handoff_data.get("file", "")
        result["comparisons"].append(handoff_result)
        handoff_issues = handoff_result["summary"]["total_issues_found"]
        total_all_issues += handoff_issues

    if support_data:
        support_result = compare_documents(wrd_data, support_data)
        support_result["comparison_label"] = "WRD vs Support"
        support_result["reference_file"] = support_data.get("file", "")
        result["comparisons"].append(support_result)
        support_issues = support_result["summary"]["total_issues_found"]
        total_all_issues += support_issues

    result["combined_summary"] = {
        "wrd_paragraphs": wrd_data.get("total_paragraphs", 0),
        "wrd_tables": wrd_data.get("total_tables", 0),
        "handoff_compared": handoff_data is not None,
        "support_compared": support_data is not None,
        "handoff_issues": handoff_issues,
        "support_issues": support_issues,
        "total_issues_all": total_all_issues,
        "compared_at": datetime.now().isoformat()
    }

    return result


# =============================================================================
# Folder scanner
# =============================================================================

def scan_and_analyze_folder(folder_path, output_fol=None):
    """Scans a folder of DOCX files, analyzes them, and writes JSON/CSV outputs."""
    p = Path(folder_path)
    if not p.exists() or not p.is_dir():
        return None, "Folder does not exist."

    results = []
    docx_files = list(p.glob("*.docx"))

    if not docx_files:
        return None, "No .docx files found in folder."

    for f in docx_files:
        try:
            analysis = analyze_docx_file(str(f))
            results.append(analysis)
        except Exception as e:
            results.append({"file": str(f), "error": str(e)})

    output_dir = Path(output_fol) if output_fol else p
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    json_path = output_dir / f"analysis_{timestamp}.json"
    with open(json_path, "w", encoding="utf-8") as jf:
        json.dump(results, jf, indent=2, default=str)

    csv_path = output_dir / f"analysis_{timestamp}.csv"
    with open(csv_path, "w", newline="", encoding="utf-8") as cf:
        writer = csv.writer(cf)
        writer.writerow(["File", "Paragraphs", "Tables", "Validation Issues"])
        for r in results:
            if "error" in r:
                writer.writerow([r["file"], "ERROR", "", r["error"]])
            else:
                writer.writerow([r["file"], r["total_paragraphs"], r["total_tables"], len(r.get("validation_issues", []))])

    return results, f"Output saved to {json_path} and {csv_path}"


# =============================================================================
# HTML Interface template - WITH FILE UPLOAD SUPPORT
# =============================================================================

HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>WRD Document Validator & Comparator</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; background: #f0f2f5; padding: 20px; }
        .container { max-width: 1000px; margin: 0 auto; }
        h1 { color: #1a1a2e; margin-bottom: 5px; text-align: center; }
        .subtitle { text-align: center; color: #666; margin-bottom: 20px; font-size: 14px; }
        .card { background: white; border-radius: 10px; padding: 25px; margin-bottom: 20px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
        .card h2 { color: #16213e; margin-bottom: 15px; font-size: 18px; }
        label { display: block; margin-bottom: 5px; font-weight: 600; color: #333; font-size: 13px; }
        .label-hint { font-weight: 400; color: #888; font-size: 12px; }
        input[type="text"] { width: 100%; padding: 10px; border: 1px solid #ddd; border-radius: 6px; margin-bottom: 10px; font-size: 14px; }
        input[type="text"]:focus { border-color: #0078d4; outline: none; box-shadow: 0 0 0 2px rgba(0,120,212,0.15); }
        button { background: #0078d4; color: white; border: none; padding: 12px 25px; border-radius: 6px; cursor: pointer; font-size: 14px; font-weight: 600; }
        button:hover { background: #006abc; }
        .results { margin-top: 20px; }
        .tab-bar { display: flex; gap: 10px; margin-bottom: 20px; }
        .tab-bar button { background: #e0e0e0; color: #333; flex: 1; }
        .tab-bar button.active { background: #0078d4; color: white; }
        .tab-content { display: none; }
        .tab-content.active { display: block; }

        /* File input group */
        .file-input-group { margin-bottom: 18px; }
        .file-input-group label { margin-bottom: 8px; }
        .input-row { display: flex; gap: 10px; align-items: center; }
        .input-row input[type="text"] { flex: 1; margin-bottom: 0; }

        /* Custom file button */
        .file-btn { display: inline-block; background: #28a745; color: white; padding: 10px 18px; border-radius: 6px; cursor: pointer; font-size: 13px; font-weight: 600; white-space: nowrap; border: none; }
        .file-btn:hover { background: #218838; }
        .file-btn input[type="file"] { display: none; }
        .file-name { font-size: 12px; color: #28a745; margin-top: 5px; font-weight: 600; min-height: 18px; }

        /* OR divider */
        .or-divider { text-align: center; color: #999; font-size: 12px; margin: 5px 0; font-weight: 600; }

        .summary-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(180px, 1fr)); gap: 15px; margin-bottom: 20px; }
        .summary-item { background: #f8f9fa; border-radius: 8px; padding: 15px; text-align: center; border: 1px solid #e9ecef; }
        .summary-item .number { font-size: 28px; font-weight: 700; color: #0078d4; }
        .summary-item .label { font-size: 12px; color: #666; margin-top: 4px; }
        .summary-item.error .number { color: #dc3545; }
        .summary-item.success .number { color: #28a745; }

        .comparison-section { margin-bottom: 25px; }
        .comparison-header { background: #e7f3ff; padding: 12px 18px; border-radius: 8px 8px 0 0; border-left: 4px solid #0078d4; font-weight: 700; font-size: 15px; color: #1a1a2e; }
        .comparison-header.support { border-left-color: #6f42c1; background: #f3e8ff; }

        .issue-card { background: #fff; border: 1px solid #e9ecef; padding: 10px 16px; margin: 1px 0; font-size: 13px; }
        .issue-card:last-child { border-radius: 0 0 8px 8px; }
        .issue-type { display: inline-block; padding: 2px 8px; border-radius: 3px; font-size: 11px; font-weight: 600; margin-right: 8px; }
        .issue-type.text { background: #dc3545; color: white; }
        .issue-type.font { background: #fd7e14; color: white; }
        .issue-type.color { background: #6f42c1; color: white; }
        .issue-type.bullet { background: #20c997; color: white; }
        .issue-type.table { background: #0dcaf0; color: #333; }
        .issue-type.missing { background: #dc3545; color: white; }
        .issue-type.spacing { background: #ffc107; color: #333; }

        pre { background: #1e1e1e; color: #d4d4d4; padding: 15px; border-radius: 6px; overflow-x: auto; font-size: 12px; max-height: 500px; overflow-y: auto; }
        .toggle-json { background: #6c757d; font-size: 12px; padding: 6px 12px; margin-bottom: 10px; }
        .no-issues { color: #28a745; font-weight: 600; padding: 15px; text-align: center; }
        .loading { text-align: center; padding: 30px; color: #666; }
    </style>
</head>
<body>
    <div class="container">
        <h1>WRD Document Validator & Comparator</h1>
        <p class="subtitle">Compare WRD against Handoff & Support reference documents</p>

        <div class="tab-bar">
            <button class="active" onclick="switchTab('compare', this)">Compare Documents</button>
            <button onclick="switchTab('analyze', this)">Analyze Folder</button>
        </div>

        <!-- Tab 1: Compare Documents -->
        <div id="tab-compare" class="tab-content active">
            <div class="card">
                <h2>Compare WRD vs Reference Documents</h2>

                <!-- WRD File -->
                <div class="file-input-group">
                    <label>WRD File (main document to validate) *</label>
                    <label class="file-btn">
                        Choose File
                        <input type="file" id="wrd_upload" accept=".docx" onchange="showFileName('wrd')">
                    </label>
                    <div class="file-name" id="wrd_filename"></div>
                    <div class="or-divider">— OR enter file path —</div>
                    <input type="text" id="wrd_path" placeholder="e.g., C:\\Documents\\WRD_main.docx">
                </div>

                <!-- Handoff File -->
                <div class="file-input-group">
                    <label>Handoff File (reference 1) <span class="label-hint">— optional if Support is provided</span></label>
                    <label class="file-btn">
                        Choose File
                        <input type="file" id="handoff_upload" accept=".docx" onchange="showFileName('handoff')">
                    </label>
                    <div class="file-name" id="handoff_filename"></div>
                    <div class="or-divider">— OR enter file path —</div>
                    <input type="text" id="handoff_path" placeholder="e.g., C:\\Documents\\Handoff_ref.docx">
                </div>

                <!-- Support File -->
                <div class="file-input-group">
                    <label>Support File (reference 2) <span class="label-hint">— optional if Handoff is provided</span></label>
                    <label class="file-btn">
                        Choose File
                        <input type="file" id="support_upload" accept=".docx" onchange="showFileName('support')">
                    </label>
                    <div class="file-name" id="support_filename"></div>
                    <div class="or-divider">— OR enter file path —</div>
                    <input type="text" id="support_path" placeholder="e.g., C:\\Documents\\Support_ref.docx">
                </div>

                <button onclick="runComparison()">Compare</button>
            </div>
        </div>

        <!-- Tab 2: Analyze Folder -->
        <div id="tab-analyze" class="tab-content">
            <div class="card">
                <h2>Analyze Folder of DOCX Files</h2>
                <label for="folder_path">Folder Path:</label>
                <input type="text" id="folder_path" placeholder="e.g., C:\\Documents\\WRD_Files">
                <label for="output_path">Output Path (optional):</label>
                <input type="text" id="output_path" placeholder="Leave blank to save in same folder">
                <button onclick="runAnalysis()">Analyze</button>
            </div>
        </div>

        <!-- Results -->
        <div class="card results" id="results-card" style="display:none;">
            <h2>Results</h2>
            <div id="results-summary"></div>
            <div id="results-details"></div>
            <button class="toggle-json" onclick="toggleJson()">Show/Hide Raw JSON</button>
            <pre id="results-json" style="display:none;"></pre>
        </div>
    </div>

    <script>
        function switchTab(tab, btn) {
            document.querySelectorAll('.tab-content').forEach(el => el.classList.remove('active'));
            document.querySelectorAll('.tab-bar button').forEach(el => el.classList.remove('active'));
            document.getElementById('tab-' + tab).classList.add('active');
            btn.classList.add('active');
        }

        function toggleJson() {
            var el = document.getElementById('results-json');
            el.style.display = el.style.display === 'none' ? 'block' : 'none';
        }

        function showFileName(prefix) {
            var input = document.getElementById(prefix + '_upload');
            var display = document.getElementById(prefix + '_filename');
            if (input.files.length > 0) {
                display.textContent = 'Selected: ' + input.files[0].name;
            } else {
                display.textContent = '';
            }
        }

        function getIssueTypeClass(type) {
            if (type.includes('TEXT')) return 'text';
            if (type.includes('FONT') || type.includes('BOLD') || type.includes('ITALIC') || type.includes('UNDERLINE') || type.includes('STRIKETHROUGH')) return 'font';
            if (type.includes('COLOR') || type.includes('HIGHLIGHT')) return 'color';
            if (type.includes('BULLET') || type.includes('INDENT') || type.includes('ALIGNMENT')) return 'bullet';
            if (type.includes('SPACE')) return 'spacing';
            if (type.includes('TABLE')) return 'table';
            if (type.includes('MISSING')) return 'missing';
            return '';
        }

        function showComparisonResults(data) {
            document.getElementById('results-card').style.display = 'block';

            var s = data.combined_summary || {};
            var summaryHtml = '<div class="summary-grid">' +
                '<div class="summary-item"><div class="number">' + (s.wrd_paragraphs || 0) + '</div><div class="label">WRD Paragraphs</div></div>' +
                '<div class="summary-item"><div class="number">' + (s.wrd_tables || 0) + '</div><div class="label">WRD Tables</div></div>' +
                '<div class="summary-item ' + (s.handoff_issues > 0 ? 'error' : 'success') + '"><div class="number">' + (s.handoff_issues || 0) + '</div><div class="label">Handoff Issues</div></div>' +
                '<div class="summary-item ' + (s.support_issues > 0 ? 'error' : 'success') + '"><div class="number">' + (s.support_issues || 0) + '</div><div class="label">Support Issues</div></div>' +
                '<div class="summary-item ' + (s.total_issues_all > 0 ? 'error' : 'success') + '"><div class="number">' + (s.total_issues_all || 0) + '</div><div class="label">Total Issues</div></div>' +
                '</div>';
            document.getElementById('results-summary').innerHTML = summaryHtml;

            var detailsHtml = '';
            (data.comparisons || []).forEach(function(comp) {
                var isSupport = comp.comparison_label.includes('Support');
                var headerClass = isSupport ? 'comparison-header support' : 'comparison-header';

                detailsHtml += '<div class="comparison-section">';
                detailsHtml += '<div class="' + headerClass + '">' + comp.comparison_label +
                    ' (' + comp.summary.total_issues_found + ' issues)</div>';

                (comp.paragraph_issues || []).forEach(function(pi) {
                    pi.issues.forEach(function(issue) {
                        detailsHtml += '<div class="issue-card">' +
                            '<span class="issue-type ' + getIssueTypeClass(issue.type) + '">' + issue.type + '</span>' +
                            ' Para ' + pi.para_id;
                        if (issue.wrd !== undefined) {
                            detailsHtml += ' | WRD: <b>' + issue.wrd + '</b> &rarr; Ref: <b>' + issue.reference + '</b>';
                        }
                        if (issue.detail) {
                            detailsHtml += ' | ' + issue.detail;
                        }
                        detailsHtml += '</div>';
                    });
                });

                (comp.table_issues || []).forEach(function(ti) {
                    ti.issues.forEach(function(issue) {
                        detailsHtml += '<div class="issue-card">' +
                            '<span class="issue-type table">' + issue.type + '</span>' +
                            ' Table ' + ti.table_id;
                        if (issue.row !== undefined) {
                            detailsHtml += ' | Row ' + issue.row + ', Col ' + issue.col;
                        }
                        if (issue.wrd !== undefined) {
                            detailsHtml += ' | WRD: <b>' + issue.wrd + '</b> &rarr; Ref: <b>' + issue.reference + '</b>';
                        }
                        if (issue.detail) {
                            detailsHtml += ' | ' + issue.detail;
                        }
                        detailsHtml += '</div>';
                    });
                });

                if ((comp.paragraph_issues || []).length === 0 && (comp.table_issues || []).length === 0) {
                    detailsHtml += '<div class="no-issues">&#10004; No issues found! Documents match.</div>';
                }

                detailsHtml += '</div>';
            });

            document.getElementById('results-details').innerHTML = detailsHtml;
            document.getElementById('results-json').textContent = JSON.stringify(data, null, 2);
        }

        function runComparison() {
            var wrdUpload = document.getElementById('wrd_upload').files[0];
            var handoffUpload = document.getElementById('handoff_upload').files[0];
            var supportUpload = document.getElementById('support_upload').files[0];

            var wrdPath = document.getElementById('wrd_path').value.trim();
            var handoffPath = document.getElementById('handoff_path').value.trim();
            var supportPath = document.getElementById('support_path').value.trim();

            var hasWrd = wrdUpload || wrdPath;
            var hasHandoff = handoffUpload || handoffPath;
            var hasSupport = supportUpload || supportPath;

            if (!hasWrd) { alert('Please provide the WRD file (upload or path).'); return; }
            if (!hasHandoff && !hasSupport) { alert('Please provide at least one reference file (Handoff or Support).'); return; }

            document.getElementById('results-card').style.display = 'block';
            document.getElementById('results-summary').innerHTML = '<div class="loading">Comparing documents...</div>';
            document.getElementById('results-details').innerHTML = '';

            // Use file upload if any file is selected
            if (wrdUpload || handoffUpload || supportUpload) {
                var formData = new FormData();
                formData.append('mode', 'upload');

                if (wrdUpload) formData.append('wrd_file', wrdUpload);
                else formData.append('wrd_path', wrdPath);

                if (handoffUpload) formData.append('handoff_file', handoffUpload);
                else if (handoffPath) formData.append('handoff_path', handoffPath);

                if (supportUpload) formData.append('support_file', supportUpload);
                else if (supportPath) formData.append('support_path', supportPath);

                fetch('/api/compare', {
                    method: 'POST',
                    body: formData
                })
                .then(r => r.json())
                .then(data => {
                    if (data.error) { alert('Error: ' + data.error); return; }
                    showComparisonResults(data);
                })
                .catch(err => alert('Error: ' + err));
            } else {
                // All paths mode
                fetch('/api/compare', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({wrd_path: wrdPath, handoff_path: handoffPath, support_path: supportPath})
                })
                .then(r => r.json())
                .then(data => {
                    if (data.error) { alert('Error: ' + data.error); return; }
                    showComparisonResults(data);
                })
                .catch(err => alert('Error: ' + err));
            }
        }

        function runAnalysis() {
            var folder = document.getElementById('folder_path').value.trim();
            var output = document.getElementById('output_path').value.trim();

            if (!folder) { alert('Please provide a folder path.'); return; }

            document.getElementById('results-card').style.display = 'block';
            document.getElementById('results-summary').innerHTML = '<div class="loading">Analyzing folder...</div>';
            document.getElementById('results-details').innerHTML = '';

            fetch('/api/analyze', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({folder_path: folder, output_path: output})
            })
            .then(r => r.json())
            .then(data => {
                if (data.error) { alert('Error: ' + data.error); return; }
                document.getElementById('results-card').style.display = 'block';
                document.getElementById('results-summary').innerHTML = '<div class="summary-grid">' +
                    '<div class="summary-item"><div class="number">' + (data.files_analyzed || 0) + '</div><div class="label">Files Analyzed</div></div>' +
                    '</div><p>' + (data.message || '') + '</p>';
                document.getElementById('results-details').innerHTML = '';
                document.getElementById('results-json').textContent = JSON.stringify(data, null, 2);
            })
            .catch(err => alert('Error: ' + err));
        }
    </script>
</body>
</html>
"""


# =============================================================================
# FLASK ROUTES
# =============================================================================

@app.route("/")
def home():
    return render_template_string(HTML_TEMPLATE)


@app.route("/api/analyze", methods=["POST"])
def run_analysis():
    data = request.get_json() or {}
    folder_path = data.get("folder_path", "").strip()
    output_path = data.get("output_path", "").strip() or None

    if not folder_path:
        return jsonify({"error": "Please provide a folder_path"}), 400

    results, message = scan_and_analyze_folder(folder_path, output_path)

    if results is None:
        return jsonify({"error": message}), 400

    return jsonify({"message": message, "files_analyzed": len(results), "results": results})


# ===== Compare endpoint - supports BOTH file upload AND file path =====
@app.route("/api/compare", methods=["POST"])
def run_comparison():
    temp_dir = None

    try:
        # Determine mode: upload or path
        is_upload = request.content_type and 'multipart/form-data' in request.content_type

        wrd_filepath = None
        handoff_filepath = None
        support_filepath = None

        if is_upload:
            # Create temp directory for uploaded files
            temp_dir = tempfile.mkdtemp(prefix="wrd_compare_")

            # WRD file
            if 'wrd_file' in request.files:
                wrd_file = request.files['wrd_file']
                if wrd_file.filename:
                    wrd_filepath = os.path.join(temp_dir, wrd_file.filename)
                    wrd_file.save(wrd_filepath)
            elif request.form.get('wrd_path', '').strip():
                wrd_filepath = request.form.get('wrd_path').strip()

            # Handoff file
            if 'handoff_file' in request.files:
                handoff_file = request.files['handoff_file']
                if handoff_file.filename:
                    handoff_filepath = os.path.join(temp_dir, handoff_file.filename)
                    handoff_file.save(handoff_filepath)
            elif request.form.get('handoff_path', '').strip():
                handoff_filepath = request.form.get('handoff_path').strip()

            # Support file
            if 'support_file' in request.files:
                support_file = request.files['support_file']
                if support_file.filename:
                    support_filepath = os.path.join(temp_dir, support_file.filename)
                    support_file.save(support_filepath)
            elif request.form.get('support_path', '').strip():
                support_filepath = request.form.get('support_path').strip()

        else:
            # JSON mode (paths only)
            data = request.get_json() or {}
            wrd_filepath = data.get("wrd_path", "").strip() or None
            handoff_filepath = data.get("handoff_path", "").strip() or None
            support_filepath = data.get("support_path", "").strip() or None

        # Validation
        if not wrd_filepath:
            return jsonify({"error": "Please provide the WRD file"}), 400

        if not handoff_filepath and not support_filepath:
            return jsonify({"error": "Please provide at least one reference file (Handoff or Support)"}), 400

        if not os.path.exists(wrd_filepath):
            return jsonify({"error": f"WRD file not found: {wrd_filepath}"}), 400

        if handoff_filepath and not os.path.exists(handoff_filepath):
            return jsonify({"error": f"Handoff file not found: {handoff_filepath}"}), 400

        if support_filepath and not os.path.exists(support_filepath):
            return jsonify({"error": f"Support file not found: {support_filepath}"}), 400

        # Analyze documents
        wrd_data = analyze_docx_file(wrd_filepath)

        handoff_data = None
        if handoff_filepath:
            handoff_data = analyze_docx_file(handoff_filepath)

        support_data = None
        if support_filepath:
            support_data = analyze_docx_file(support_filepath)

        # Run comparison
        comparison = compare_three_documents(wrd_data, handoff_data, support_data)

        return jsonify(comparison)

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        # Clean up temp files
        if temp_dir and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)


# =============================================================================
# MAIN
# =============================================================================

if __name__ == "__main__":
    print("-" * 60)
    print("  WRD Document Validator & Comparator")
    print("  ====================================")
    print("  Features:")
    print("    - File upload support (Choose File)")
    print("    - Font color detection")
    print("    - Paragraph alignment check")
    print("    - Indentation validation")
    print("    - Bullet / numbering detection")
    print("    - Bullet level consistency")
    print("    - Table comparison")
    print("    - Spacing validation")
    print("    - Standards validation")
    print("    - 3-file comparison: WRD vs Handoff + WRD vs Support")
    print("")
    print("  Requirements: pip install flask python-docx lxml")
    print("-" * 60)
    app.run(host="127.0.0.1", port=5000, debug=True)
