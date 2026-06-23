# -*- coding: utf-8 -*-
"""
DOCX Formatting Validator - Python Local Server Code
===================================================
A Flask-based application for scanning and validating the formatting of .docx documents.
Performs XML-level run parsing and paragraph-level validation against style conventions.

Requires:
    pip install flask python-docx lxml

Usage:
    python app.py
    Then direct your browser to http://127.0.0.1:5000
"""

import os
import json
import csv
import re
from pathlib import Path
from datetime import datetime
from flask import Flask, request, jsonify, render_template_string
from docx.shared import RGBColor


# Try to import python-docx
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Warning: 'python-docx' library is not installed. Install with 'pip install python-docx'")

app = Flask(__name__)

# Predefined standard rules as requested in the prompt
DEFAULT_STANDARDS = [
    {
        "id": "heading1",
        "name": "Level 1 Heading",
        "fontName": "Human Black Condensed",
        "fontSize": 20.0,
        "bold": True,
        "alignment": "LEFT",
        "indent": 0.0
    },
    {
        "id": "heading2",
        "name": "Level 2 Heading",
        "fontName": "Human Black Condensed",
        "fontSize": 11.5,
        "bold": True,
        "alignment": "LEFT",
        "indent": 0.0
    },
    {
        "id": "heading3",
        "name": "Level 3 Heading",
        "fontName": "Human Black Condensed",
        "fontSize": 10.5,
        "bold": True,
        "alignment": "LEFT",
        "indent": 0.0
    },
    {
        "id": "normal",
        "name": "Normal Text",
        "fontName": "Human Light Condensed",
        "fontSize": 10.5,
        "bold": False,
        "alignment": "LEFT",
        "indent": 0.0
    },
    {
        "id": "boldText",
        "name": "Text Level Bold",
        "fontName": "Human Bold Condensed",
        "fontSize": 10.5,
        "bold": True,
        "alignment": "LEFT",
        "indent": 0.0
    },
    {
        "id": "formName",
        "name": "Form Name",
        "fontName": "Human Light Condensed",
        "fontSize": 7.0,
        "bold": False,
        "alignment": "LEFT",
        "indent": 0.0
    }
]

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

def get_doc_defaults(doc):
    """Parse styles.xml to get default document properties."""
    defaults = {"fontName": None, "fontSize": None, "bold": None}
    try:
        styles_element = doc.styles.element
        doc_def = styles_element.find(f"{W_NS}docDefaults")
        if doc_def is not None:
            rpr_default = doc_def.find(f"{W_NS}rPrDefault")
            if rpr_default is not None:
                rpr = rpr_default.find(f"{W_NS}rPr")
                if rpr is not None:
                    rfonts = rpr.find(f"{W_NS}rFonts")
                    if rfonts is not None:
                        for attr in ["ascii", "hAnsi", "cs"]:
                            val = rfonts.get(f"{W_NS}{attr}")
                            if val:
                                defaults["fontName"] = val
                                break
                    sz = rpr.find(f"{W_NS}sz")
                    if sz is not None:
                        hp = sz.get(f"{W_NS}val")
                        if hp:
                            defaults["fontSize"] = int(hp) / 2.0
                    b_elem = rpr.find(f"{W_NS}b")
                    if b_elem is not None:
                        b_val = b_elem.get(f"{W_NS}val")
                        defaults["bold"] = b_val is None or b_val.lower() in ("true", "1")
    except Exception:
        pass
    return defaults


def get_xml_run_props(r_elem):
    """Extract formatting values from a single run elements."""
    props = {"fontName": None, "fontSize": None, "bold": None, "runStyleId": None}
    rpr = r_elem.find(f"{W_NS}rPr")
    if rpr is None:
        return props

    rfonts = rpr.find(f"{W_NS}rFonts")
    if rfonts is not None:
        for attr in ["ascii", "hAnsi", "cs"]:
            val = rfonts.get(f"{W_NS}{attr}")
            if val:
                props["fontName"] = val
                break

    sz = rpr.find(f"{W_NS}sz")
    if sz is not None:
        val = sz.get(f"{W_NS}val")
        if val:
            props["fontSize"] = int(val) / 2.0

    b_elem = rpr.find(f"{W_NS}b")
    if b_elem is not None:
        val = b_elem.get(f"{W_NS}val")
        props["bold"] = val is None or val.lower() in ("true", "1")

    r_style = rpr.find(f"{W_NS}rStyle")
    if r_style is not None:
        props["runStyleId"] = r_style.get(f"{W_NS}val")

    return props


def walk_style_chain(style, key_func, max_depth=10):
    """Traverse style hierarchy (based_on) to resolve formatting."""
    current = style
    depth = 0
    while current and depth < max_depth:
        val = key_func(current)
        if val is not None:
            return val
        try:
            current = current.base_style
        except Exception:
            break
        depth += 1
    return None


def resolve_properties(run_props, para, doc_defaults, doc_styles):
    """Resolve final font properties based on XML cascade."""
    font_name = run_props["fontName"]
    font_size = run_props["fontSize"]
    bold = run_props["bold"]

    # Try run style
    if run_props["runStyleId"] and run_props["runStyleId"] in doc_styles:
        rstyle = doc_styles[run_props["runStyleId"]]
        if font_name is None:
            font_name = walk_style_chain(rstyle, lambda s: s.font.name if s.font else None)
        if font_size is None:
            font_size = walk_style_chain(rstyle, lambda s: s.font.size.pt if s.font and s.font.size else None)
        if bold is None:
            bold = walk_style_chain(rstyle, lambda s: s.font.bold if s.font else None)

    # Try paragraph style
    try:
        pstyle = para.style
        if pstyle:
            if font_name is None:
                font_name = walk_style_chain(pstyle, lambda s: s.font.name if s.font else None)
            if font_size is None:
                font_size = walk_style_chain(pstyle, lambda s: s.font.size.pt if s.font and s.font.size else None)
            if bold is None:
                bold = walk_style_chain(pstyle, lambda s: s.font.bold if s.font else None)
    except Exception:
        pass

    # Fallback to document defaults
    if font_name is None:
        font_name = doc_defaults["fontName"]
    if font_size is None:
        font_size = doc_defaults["fontSize"]
    if bold is None:
        bold = doc_defaults["bold"]

    return font_name or "Calibri", font_size or 11.0, bold or False


def parse_paragraph_xml_runs(para, doc_defaults, doc_styles):
    """Parser to find nested run blocks in paragraph and gather words."""
    p_elem = para._element
    runs_data = []
    
    for r_elem in p_elem.iter(f"{W_NS}r"):
        # Combine nested text tags
        text_parts = []
        for child in r_elem:
            if child.tag == f"{W_NS}t" and child.text:
                text_parts.append(child.text)
            elif child.tag == f"{W_NS}tab":
                text_parts.append("\t")
            elif child.tag == f"{W_NS}br":
                text_parts.append("\n")
        
        run_text = "".join(text_parts)
        if not run_text:
            continue

        run_xml_props = get_xml_run_props(r_elem)
        r_font, r_size, r_bold = resolve_properties(run_xml_props, para, doc_defaults, doc_styles)
        
        runs_data.append({
            "text": run_text,
            "fontName": r_font,
            "fontSize": r_size,
            "bold": r_bold
        })
    return runs_data


def analyze_docx_file(filepath, standards=DEFAULT_STANDARDS):
    """Extract paragraphs and run metrics with style validations."""
    doc = Document(str(filepath))
    doc_defaults = get_doc_defaults(doc)
    doc_styles = {s.style_id: s for s in doc.styles if s.style_id}
    
    paragraphs_results = []
    
    for idx, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue # skip blank lines
            
        runs = parse_paragraph_xml_runs(para, doc_defaults, doc_styles)
        
        # Determine dominant styles
        font_lens = {}
        size_lens = {}
        bold_lens = {True: 0, False: 0}
        
        for r in runs:
            rlen = len(r["text"])
            font_lens[r["fontName"]] = font_lens.get(r["fontName"], 0) + rlen
            size_lens[r["fontSize"]] = size_lens.get(r["fontSize"], 0) + rlen
            bold_lens[r["bold"]] = bold_lens.get(r["bold"], 0) + rlen
            
        dominant_font = max(font_lens, key=font_lens.get) if font_lens else "Calibri"
        dominant_size = max(size_lens, key=size_lens.get) if size_lens else 11.0
        dominant_bold = bold_lens[True] >= bold_lens[False] if runs else False
        
        # Intox alignment mapping
        align_str = "LEFT"
        if para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            align_str = "RIGHT"
        elif para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align_str = "CENTER"
        elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            align_str = "JUSTIFY"
            
        # Left indent mapping (dxa -> points)
        indent_pt = 0
        if para.paragraph_format.left_indent:
            indent_pt = round(para.paragraph_format.left_indent.pt, 2)
            
        style_name = para.style.name if para.style else "Normal"
        
        # Run validation matching Closest Level
        matched_rule = None
        min_diff = 9999
        for rule in standards:
            diff = abs(dominant_size - rule["fontSize"]) * 10
            if dominant_font.lower() != rule["fontName"].lower():
                diff += 5
            if rule["bold"] is not None and dominant_bold != rule["bold"]:
                diff += 3
            if style_name.lower().find("heading") != -1 and rule["id"].startswith("heading"):
                diff -= 2 # style correlation bonus
            if diff < min_diff:
                min_diff = diff
                matched_rule = rule
                
        issues = []
        validation_status = "Pass"
        matched_level = "Unknown Standard"
        matched_level_id = "unmatched"
        
        if matched_rule:
            matched_level = matched_rule["name"]
            matched_level_id = matched_rule["id"]
            
            # Font validation
            if dominant_font.lower() != matched_rule["fontName"].lower():
                issues.append({
                    "type": "fontName",
                    "severity": "error",
                    "message": f"Incorrect Font for {matched_rule['name']}. Expected '{matched_rule['fontName']}', but found '{dominant_font}'."
                })
            # Size validation
            if dominant_size != matched_rule["fontSize"]:
                issues.append({
                    "type": "fontSize",
                    "severity": "error",
                    "message": f"Incorrect Font Size for {matched_rule['name']}. Expected {matched_rule['fontSize']}pt, but found {dominant_size}pt."
                })
            # Bold validation
            if matched_rule["bold"] is not None and dominant_bold != matched_rule["bold"]:
                issues.append({
                    "type": "bold",
                    "severity": "error",
                    "message": f"Incorrect Bold formatting for {matched_rule['name']}. Expected {'Bold' if matched_rule['bold'] else 'Normal'}."
                })
            # Alignment validation
            if matched_rule["alignment"] and align_str != matched_rule["alignment"]:
                issues.append({
                    "type": "alignment",
                    "severity": "warning",
                    "message": f"Inconsistent alignment. Expected '{matched_rule['alignment']}', found '{align_str}'."
                })
            # Indentation check
            if matched_rule["indent"] is not None and indent_pt != matched_rule["indent"]:
                # Flag if mismatch is significant
                if abs(indent_pt - matched_rule["indent"]) > 1.0:
                    issues.append({
                        "type": "indent",
                        "severity": "warning",
                        "message": f"Indentation issue. Expected '{matched_rule['indent']}pt', found '{indent_pt}pt'."
                    })
        else:
            issues.append({
                "type": "unmatched",
                "severity": "error",
                "message": f"Paragraph formatting matches no known standards."
            })
            
        any_errors = any(i["severity"] == "error" for i in issues)
        if any_errors or not matched_rule:
            validation_status = "Fail"
            
        paragraphs_results.append({
            "paragraphNumber": idx + 1,
            "text": text,
            "styleName": style_name,
            "alignment": align_str,
            "indent": indent_pt,
            "dominantFont": dominant_font,
            "dominantSize": dominant_size,
            "dominantBold": dominant_bold,
            "matchedLevelId": matched_level_id,
            "matchedLevelName": matched_level,
            "validationStatus": validation_status,
            "issues": issues,
            "runs": runs
        })
        
    return paragraphs_results


def scan_and_analyze_folder(folder_path, output_fol=None):
    """Scans whole folder of DOCX files, analyzes them, and writes local JSON/CSV outputs."""
    p = Path(folder_path)
    if not p.exists() or not p.is_dir():
        return None, "Folder does not exist."
        
    files = list(p.glob("*.docx"))
    if not files:
        return None, "No .docx files found in this folder."
        
    all_reports = []
    
    for f in files:
        try:
            p_res = analyze_docx_file(f)
            
            # Compute stats
            total = len(p_res)
            passed = sum(1 for x in p_res if x["validationStatus"] == "Pass")
            failed = total - passed
            issues_cnt = sum(len(x["issues"]) for x in p_res)
            
            summary = {
                "totalParagraphs": total,
                "passed": passed,
                "failed": failed,
                "issuesCount": issues_cnt
            }
            
            all_reports.append({
                "fileName": f.name,
                "fileSize": f.stat().st_size,
                "filePath": str(f),
                "paragraphs": p_res,
                "summary": summary
            })
            
        except Exception as e:
            all_reports.append({
                "fileName": f.name,
                "filePath": str(f),
                "paragraphs": [],
                "error": str(e),
                "summary": {"totalParagraphs": 0, "passed": 0, "failed": 0, "issuesCount": 0}
            })
            
    # Save output to specified folder if provided
    if output_fol:
        out_p = Path(output_fol)
        out_p.mkdir(parents=True, exist_ok=True)
        
        # Write clean combined JSON
        json_path = out_p / "formatting_analysis_report.json"
        with open(json_path, "w", encoding="utf-8") as jf:
            json.dump(all_reports, jf, ensure_ascii=False, indent=2)
            
        # Write clean flat CSV of results
        csv_path = out_p / "formatting_analysis_report.csv"
        with open(csv_path, "w", encoding="utf-8", newline="") as cf:
            writer = csv.writer(cf)
            writer.writerow([
                "File Name", "Paragraph Num", "Paragraph Text", 
                "Font Name", "Font Size", "Bold", "Alignment", "Indent (pt)",
                "Expected Level", "Status", "Issues Highlight"
            ])
            for report in all_reports:
                for p_obj in report.get("paragraphs", []):
                    issue_list = "; ".join([i["message"] for i in p_obj["issues"]])
                    writer.writerow([
                        report["fileName"],
                        p_obj["paragraphNumber"],
                        p_obj["text"][:120].replace("\n", " ") + ("..." if len(p_obj["text"]) > 120 else ""),
                        p_obj["dominantFont"],
                        p_obj["dominantSize"],
                        "True" if p_obj["dominantBold"] else "False",
                        p_obj["alignment"],
                        p_obj["indent"],
                        p_obj["matchedLevelName"],
                        p_obj["validationStatus"],
                        issue_list
                    ])
                    
    return all_reports, None

# HTML Interface template for Flask
HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>DocuLint Compliance Assessment Local Board</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
    <script src="https://cdn.tailwindcss.com"></script>
    <script>
        tailwind.config = {
            theme: {
                extend: {
                    fontFamily: {
                        sans: ['Inter', 'ui-sans-serif', 'system-ui', 'sans-serif'],
                        mono: ['JetBrains Mono', 'ui-monospace', 'SFMono-Regular', 'monospace'],
                    }
                }
            }
        }
    </script>
    <style>
        body { font-family: 'Inter', sans-serif; }
        .gradient-bg { background: linear-gradient(135deg, #4f46e5 0%, #3730a3 100%); }
    </style>
</head>
<body class="bg-slate-50 text-slate-800 min-h-screen flex flex-col font-sans antialiased">
    <!-- Top Nav -->
    <header class="bg-white border-b border-slate-200 sticky top-0 z-50 shadow-xs">
        <div class="max-w-7xl mx-auto px-6 py-4 flex items-center justify-between">
            <div class="flex items-center space-x-3">
                <div class="bg-indigo-600 p-2 rounded-lg text-white">
                    <svg class="w-6 h-6" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                        <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9 12h6m-6 4h6m2 5H7a2 2 0 01-2-2V5a2 2 0 012-2h5.586a1 1 0 01.707.293l5.414 5.414a1 1 0 01.293.707V19a2 2 0 01-2 2z" />
                    </svg>
                </div>
                <div>
                    <h1 class="text-lg font-bold tracking-tight text-slate-900 leading-none">DocuLint <span class="text-indigo-600 font-semibold">Local</span></h1>
                    <p class="text-[10px] text-slate-400 mt-1 font-mono tracking-tight font-medium">STRUCTURAL FORMATTING STANDARDS VALIDATOR</p>
                </div>
            </div>
            <div class="flex items-center gap-3">
                <span class="inline-flex items-center px-2.5 py-0.5 rounded-full text-xs font-semibold bg-emerald-50 text-emerald-700 border border-emerald-100">
                    <span class="w-1.5 h-1.5 mr-1.5 rounded-full bg-emerald-500 animate-pulse"></span>
                    Local Python Server Active
                </span>
            </div>
        </div>
    </header>

    <main class="max-w-7xl mx-auto px-6 py-8 flex-1 w-full space-y-6">
        
        <!-- Welcome grid (Form and Standards definition panel) -->
        <div class="grid grid-cols-1 lg:grid-cols-3 gap-6">
            
            <!-- Left 2 columns: Scan config form -->
            <div class="lg:col-span-2 bg-white rounded-2xl border border-slate-200 p-6 shadow-sm flex flex-col justify-between">
                <div>
                    <h2 class="text-sm font-bold text-slate-800 font-display flex items-center gap-2">
                        <svg class="w-4 h-4 text-indigo-500" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M19 11H5m14 0a2 2 0 012 2v6a2 2 0 01-2 2H5a2 2 0 01-2-2v-6a2 2 0 012-2m14 0V9a2 2 0 00-2-2M5 11V9a2 2 0 012-2m0 0V5a2 2 0 012-2h6a2 2 0 012 2v2M7 7h10" /></svg>
                        Target Scanning Directories
                    </h2>
                    <p class="text-xs text-slate-400 mt-1 mb-6 leading-relaxed">
                        Specify the local folder path containing your <code class="bg-slate-100 text-slate-700 px-1 py-0.5 rounded font-mono">.docx</code> documents. DocuLint will traverse and analyze them at XML level directly from your filesystem.
                    </p>
                    
                    <form id="analyzeForm" class="space-y-4">
                        <div class="grid grid-cols-1 gap-4">
                            <div>
                                <label class="block text-[11px] font-bold text-slate-500 uppercase tracking-wider mb-1.5">Local Scanning Directory (Absolute Path) <span class="text-rose-500">*</span></label>
                                <div class="relative">
                                    <input type="text" id="folderPath" class="w-full px-4 py-2.5 bg-slate-50 hover:bg-slate-100/50 focus:bg-white border border-slate-200 rounded-xl outline-none focus:border-indigo-600 focus:ring-2 focus:ring-indigo-150 transition-all font-mono text-xs text-slate-700" placeholder="e.g. C:/Users/YourName/Documents/CheckFolder" required>
                                </div>
                            </div>
                            
                            <div>
                                <label class="block text-[11px] font-bold text-slate-500 uppercase tracking-wider mb-1.5">Output Report Directory (Optional)</label>
                                <div class="relative">
                                    <input type="text" id="outputPath" class="w-full px-4 py-2.5 bg-slate-50 hover:bg-slate-100/50 focus:bg-white border border-slate-200 rounded-xl outline-none focus:border-indigo-600 focus:ring-2 focus:ring-indigo-150 transition-all font-mono text-xs text-slate-700" placeholder="e.g. ./output_reports (Generates JSON & CSV on disk)">
                                </div>
                            </div>
                        </div>

                        <div class="pt-4 flex items-center justify-between border-t border-slate-100 mt-4">
                            <span class="text-[10px] text-slate-400 font-medium">Uses standard default style hierarchies check matrix</span>
                            <button type="submit" id="btn_submit" class="px-5 py-2.5 bg-gradient-to-r from-indigo-600 to-indigo-700 hover:from-indigo-700 hover:to-indigo-800 text-white font-semibold text-xs rounded-xl shadow-md tracking-wide active:scale-95 transition-all outline-none cursor-pointer flex items-center gap-1.5">
                                <svg class="w-4 h-4" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M14.752 11.168l-3.197-2.132A1 1 0 0010 9.87v4.263a1 1 0 001.555.832l3.197-2.132a1 1 0 000-1.664z" /><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M21 12a9 9 0 11-18 0 9 9 0 0118 0z" /></svg>
                                Run Formatting Scans
                            </button>
                        </div>
                    </form>
                </div>
            </div>

            <!-- Right 1 column: Validation Standards Applied list -->
            <div class="bg-white rounded-2xl border border-slate-200 p-6 shadow-sm">
                <h3 class="text-sm font-bold text-slate-800 font-display flex items-center gap-2">
                    <svg class="w-4 h-4 text-indigo-500" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M10.325 4.317c.426-1.756 2.924-1.756 3.35 0a1.724 1.724 0 002.573 1.066c1.543-.94 3.31.826 2.37 2.37a1.724 1.724 0 001.065 2.572c1.756.426 1.756 2.924 0 3.35a1.724 1.724 0 00-1.066 2.573c.94 1.543-.826 3.31-2.37 2.37a1.724 1.724 0 00-2.572 1.065c-.426 1.756-2.924 1.756-3.35 0a1.724 1.724 0 00-2.573-1.066-1.543.94-3.31-.826-2.37-2.37a1.724 1.724 0 00-1.065-2.572c-1.756-.426-1.756-2.924 0-3.35a1.724 1.724 0 001.066-2.573c-.94-1.543.826-3.31 2.37-2.37.996.608 2.296.07 2.572-1.065z" /><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M15 12a3 3 0 11-6 0 3 3 0 016 0z" /></svg>
                    Enforced Rule Standards
                </h3>
                <p class="text-xs text-slate-400 mt-1 mb-4 leading-relaxed">
                    Documents are automatically validated against the standard configuration matrices:
                </p>

                <div class="space-y-2 max-h-[220px] overflow-y-auto pr-1">
                    <div class="div-standard-item p-3 rounded-xl border border-slate-200 bg-slate-50/50 flex flex-col gap-1.5 text-xs">
                        <div class="flex items-center justify-between font-semibold text-slate-800">
                            <span>Heading 1</span>
                            <span class="text-[10px] font-mono bg-indigo-50 text-indigo-600 px-1.5 py-0.5 rounded">20 pt</span>
                        </div>
                        <div class="text-[10px] text-slate-400 font-mono flex flex-wrap gap-x-2">
                            <span>Font: <strong class="text-slate-600">Human Black Condensed</strong></span>
                            <span>•</span>
                            <span>Bold: <strong class="text-slate-600">True</strong></span>
                            <span>•</span>
                            <span>Align: <strong class="text-slate-600">LEFT</strong></span>
                        </div>
                    </div>

                    <div class="div-standard-item p-3 rounded-xl border border-slate-200 bg-slate-50/50 flex flex-col gap-1.5 text-xs">
                        <div class="flex items-center justify-between font-semibold text-slate-800">
                            <span>Heading 2</span>
                            <span class="text-[10px] font-mono bg-indigo-50 text-indigo-600 px-1.5 py-0.5 rounded">11.5 pt</span>
                        </div>
                        <div class="text-[10px] text-slate-400 font-mono flex flex-wrap gap-x-2">
                            <span>Font: <strong class="text-slate-600">Human Black Condensed</strong></span>
                            <span>•</span>
                            <span>Bold: <strong class="text-slate-600">True</strong></span>
                            <span>•</span>
                            <span>Align: <strong class="text-slate-600">LEFT</strong></span>
                        </div>
                    </div>

                    <div class="div-standard-item p-3 rounded-xl border border-slate-200 bg-slate-50/50 flex flex-col gap-1.5 text-xs">
                        <div class="flex items-center justify-between font-semibold text-slate-800">
                            <span>Normal Paragraph Body</span>
                            <span class="text-[10px] font-mono bg-indigo-50 text-indigo-600 px-1.5 py-0.5 rounded">10.5 pt</span>
                        </div>
                        <div class="text-[10px] text-slate-400 font-mono flex flex-wrap gap-x-2">
                            <span>Font: <strong class="text-slate-600">Human Light Condensed</strong></span>
                            <span>•</span>
                            <span>Bold: <strong class="text-slate-600">False</strong></span>
                            <span>•</span>
                            <span>Align: <strong class="text-slate-600">LEFT</strong></span>
                        </div>
                    </div>

                    <div class="div-standard-item p-3 rounded-xl border border-slate-200 bg-slate-50/50 flex flex-col gap-1.5 text-xs">
                        <div class="flex items-center justify-between font-semibold text-slate-800">
                            <span>Form Name</span>
                            <span class="text-[10px] font-mono bg-indigo-50 text-indigo-600 px-1.5 py-0.5 rounded">7 pt</span>
                        </div>
                        <div class="text-[10px] text-slate-400 font-mono flex flex-wrap gap-x-2">
                            <span>Font: <strong class="text-slate-600">Human Light Condensed</strong></span>
                            <span>•</span>
                            <span>Bold: <strong class="text-slate-600">False</strong></span>
                            <span>•</span>
                            <span>Align: <strong class="text-slate-600">LEFT</strong></span>
                        </div>
                    </div>
                </div>
            </div>
        </div>

        <!-- Dynamic Results Segment -->
        <div id="resultsArea" class="hidden space-y-6">
            
            <!-- Dashboard Bento Summary Card with stats cards inside -->
            <div class="bg-indigo-900 text-white rounded-2xl p-6 shadow-md gradient-bg animate-fade-in">
                <div class="flex flex-col sm:flex-row items-start sm:items-center justify-between gap-4 mb-6 pb-6 border-b border-white/10">
                    <div>
                        <h3 class="text-lg font-bold font-display">Formatting Compliance Dashboard</h3>
                        <p class="text-xs text-indigo-200 mt-1">Real-time local metrics from parsed active standard layouts</p>
                    </div>
                </div>
                
                <div class="grid grid-cols-2 md:grid-cols-4 gap-4" id="statsGrid">
                    <!-- Dynamic stats loaded by JS -->
                </div>
            </div>

            <!-- Controls bar with Search and Filters -->
            <div class="bg-white rounded-xl border border-slate-200 p-4 flex flex-col sm:flex-row items-center gap-4 justify-between shadow-xs">
                <div class="relative w-full sm:max-w-xs">
                    <input 
                        type="text" 
                        id="searchQuery" 
                        placeholder="Search text contents..." 
                        class="w-full pl-9 pr-4 py-2 bg-slate-50 border border-slate-200 rounded-lg text-xs outline-none focus:bg-white focus:border-indigo-500 transition"
                    >
                    <svg class="w-4 h-4 text-slate-400 absolute left-3 top-2.5" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M21 21l-6-6m2-5a7 7 0 11-14 0 7 7 0 0114 0z" /></svg>
                </div>

                <div class="flex items-center gap-2 w-full sm:w-auto shrink-0 justify-end">
                    <span class="text-xs text-slate-500 font-medium font-sans">Filter check:</span>
                    <button id="filterAll" class="px-3 py-1.5 bg-indigo-50 text-indigo-600 border border-indigo-150 text-xs font-semibold rounded-lg transition active:scale-95 cursor-pointer">
                        All rows
                    </button>
                    <button id="filterFailed" class="px-3 py-1.5 bg-rose-50 hover:bg-rose-100 text-rose-600 text-xs font-semibold rounded-lg transition border border-rose-100 active:scale-95 cursor-pointer flex items-center gap-1.5">
                        <span class="w-1.5 h-1.5 bg-rose-500 rounded-full"></span>
                        Show failures only
                    </button>
                </div>
            </div>

            <!-- List of Verified Files -->
            <div id="filesContainer" class="space-y-6">
                <!-- Dynamic File Cards -->
            </div>
            
        </div>
    </main>

    <footer class="bg-white border-t border-slate-200 py-6 mt-12 text-center text-xs text-slate-400 font-mono">
        📄 Generated securely locally. Powered by DocuLint local Flask server wrapper.
    </footer>

    <script>
        let globalReportsData = [];
        let curFilterMode = 'all'; // 'all' or 'failed'
        let curSearchQuery = '';

        document.getElementById('analyzeForm').addEventListener('submit', async (e) => {
            e.preventDefault();
            const folder_path = document.getElementById('folderPath').value;
            const output_path = document.getElementById('outputPath').value;
            
            const btn = document.getElementById('btn_submit');
            btn.disabled = true;
            const oldHtml = btn.innerHTML;
            btn.innerHTML = `
                <svg class="animate-spin -ml-1 mr-2 h-4 w-4 text-white" fill="none" viewBox="0 0 24 24">
                    <circle class="opacity-25" cx="12" cy="12" r="10" stroke="currentColor" stroke-width="4"></circle>
                    <path class="opacity-75" fill="currentColor" d="M4 12a8 8 0 018-8V0C5.373 0 0 5.373 0 12h4zm2 5.291A7.962 7.962 0 014 12H0c0 3.042 1.135 5.824 3 7.938l3-2.647z"></path>
                </svg>
                Extracting DOCX...
            `;
            
            try {
                const res = await fetch('/api/analyze', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ folder_path, output_path })
                });
                const data = await res.json();
                
                if (data.error) {
                    alert('Error: ' + data.error);
                } else {
                    globalReportsData = data.reports;
                    renderUI();
                }
            } catch (err) {
                alert('Analysis request failed on local python server. Make sure folder paths are absolute and exist on disk.');
            } finally {
                btn.disabled = false;
                btn.innerHTML = oldHtml;
            }
        });

        // Search inputs listeners
        document.getElementById('searchQuery').addEventListener('input', (e) => {
            curSearchQuery = e.target.value.toLowerCase().trim();
            renderUI();
        });

        document.getElementById('filterAll').addEventListener('click', () => {
            curFilterMode = 'all';
            document.getElementById('filterAll').className = 'px-3 py-1.5 bg-indigo-50 text-indigo-600 border border-indigo-150 text-xs font-semibold rounded-lg transition active:scale-95 cursor-pointer';
            document.getElementById('filterFailed').className = 'px-3 py-1.5 bg-rose-50 hover:bg-rose-100 text-rose-600 text-xs font-semibold rounded-lg transition border border-rose-100 active:scale-95 cursor-pointer flex items-center gap-1.5';
            renderUI();
        });

        document.getElementById('filterFailed').addEventListener('click', () => {
            curFilterMode = 'failed';
            document.getElementById('filterFailed').className = 'px-3 py-1.5 bg-rose-100 text-rose-700 border border-rose-200 text-xs font-semibold rounded-lg transition active:scale-95 cursor-pointer flex items-center gap-1.5';
            document.getElementById('filterAll').className = 'px-3 py-1.5 bg-slate-50 hover:bg-slate-100 text-slate-600 text-xs font-semibold rounded-lg transition active:scale-95 cursor-pointer';
            renderUI();
        });

        function renderUI() {
            const resultsArea = document.getElementById('resultsArea');
            resultsArea.classList.remove('hidden');

            // Quick calculations for total stats
            let totalFiles = globalReportsData.length;
            let totalParas = 0;
            let passedParas = 0;
            let failedParas = 0;
            let totalIssues = 0;

            globalReportsData.forEach(r => {
                totalParas += r.summary.totalParagraphs;
                passedParas += r.summary.passed;
                failedParas += r.summary.failed;
                totalIssues += r.summary.issuesCount;
            });

            // Compliance percentage
            const scoreRate = totalParas ? ((passedParas / totalParas) * 100).toFixed(1) : '100.0';

            // Populate dashboard stats
            const statsGrid = document.getElementById('statsGrid');
            statsGrid.innerHTML = `
                <div class="bg-white/10 backdrop-blur-xs p-4 rounded-xl border border-white/10">
                    <p class="text-[10px] text-indigo-200 font-bold uppercase tracking-wider">Scanned Files</p>
                    <p class="text-2xl font-bold font-display mt-0.5">${totalFiles}</p>
                </div>
                <div class="bg-white/10 backdrop-blur-xs p-4 rounded-xl border border-white/10">
                    <p class="text-[10px] text-indigo-200 font-bold uppercase tracking-wider">Total Checked</p>
                    <p class="text-2xl font-bold font-display mt-0.5">${totalParas} lines</p>
                </div>
                <div class="bg-white/10 backdrop-blur-xs p-4 rounded-xl border border-white/10">
                    <p class="text-[10px] text-indigo-200 font-bold uppercase tracking-wider">Passed / Failed</p>
                    <p class="text-2xl font-bold mt-0.5 font-display flex items-center gap-1.5">
                        <span class="text-emerald-300 font-bold">${passedParas}</span>
                        <span class="text-indigo-300 font-normal">/</span>
                        <span class="${failedParas > 0 ? 'text-rose-300' : 'text-indigo-200'} font-bold">${failedParas}</span>
                    </p>
                </div>
                <div class="bg-white/10 backdrop-blur-xs p-4 rounded-xl border border-white/10">
                    <p class="text-[10px] text-indigo-200 font-bold uppercase tracking-wider">Compliance Rate</p>
                    <p class="text-2xl font-bold text-emerald-300 mt-0.5 font-display">${scoreRate}%</p>
                </div>
            `;

            // Populate file listings
            const container = document.getElementById('filesContainer');
            container.innerHTML = '';

            let renderedFilesCount = 0;

            globalReportsData.forEach((f, idx) => {
                let fileRowsHtml = '';
                let matchingParagraphsCount = 0;

                f.paragraphs.forEach(p => {
                    // Filter matching query or failures
                    if (curFilterMode === 'failed' && p.validationStatus === 'Pass') return;
                    if (curSearchQuery && !p.text.toLowerCase().includes(curSearchQuery)) return;

                    matchingParagraphsCount++;

                    const statusClass = p.validationStatus === 'Pass' 
                        ? 'bg-transparent text-slate-700' 
                        : 'bg-rose-50/40 border-l-4 border-rose-500';
                        
                    const statusBadge = p.validationStatus === 'Pass'
                        ? '<span class="text-emerald-600 font-semibold px-2 py-0.5 rounded text-[10px] bg-emerald-50 border border-emerald-100">PASS</span>'
                        : '<span class="text-rose-600 font-semibold px-2 py-0.5 rounded text-[10px] bg-rose-50 border border-rose-100 animate-pulse">FAIL</span>';

                    const issueTexts = p.issues.map(issue => `
                        <div class="text-xs text-rose-600 flex items-start gap-1 font-sans mt-0.5">
                            <span class="text-rose-500 font-semibold">•</span>
                            <span>${issue.message}</span>
                        </div>
                    `).join('');

                    fileRowsHtml += `
                        <tr class="${statusClass} hover:bg-slate-50/70 border-b border-slate-100 transition">
                            <td class="px-5 py-3 text-[10px] font-mono text-slate-400 text-center">${p.paragraphNumber}</td>
                            <td class="px-5 py-3 max-w-md">
                                <p class="text-xs font-semibold text-slate-800 leading-relaxed">${escapeHtml(p.text)}</p>
                                <div class="mt-1.5 text-[9px] text-slate-400 space-x-2">
                                    <span>Style: <strong class="text-slate-600 font-mono">${p.styleName}</strong></span>
                                    <span>•</span>
                                    <span>Align: <strong class="text-slate-600">${p.alignment}</strong></span>
                                    <span>•</span>
                                    <span>Indent: <strong class="text-slate-600">${p.indent}pt</strong></span>
                                </div>
                            </td>
                            <td class="px-5 py-3 text-[10px] space-y-0.5 font-mono">
                                <div class="text-slate-500">Font: <span class="text-slate-800 font-semibold">${p.dominantFont}</span></div>
                                <div class="text-slate-500">Size: <span class="bg-slate-100 text-slate-705 px-1 py-0.5 rounded">${p.dominantSize}pt</span></div>
                                <div class="text-slate-500">Bold: <span class="text-slate-600">${p.dominantBold ? 'True' : 'False'}</span></div>
                            </td>
                            <td class="px-5 py-3 text-[10px] font-bold text-indigo-600 font-mono">${p.matchedLevelName}</td>
                            <td class="px-5 py-3 text-center">${statusBadge}</td>
                            <td class="px-5 py-3 max-w-xs">${issueTexts || '<span class="text-[10px] text-emerald-600 font-semibold flex items-center gap-1"><span class="w-1.5 h-1.5 rounded-full bg-emerald-500"></span>Standard Compliant</span>'}</td>
                        </tr>
                    `;
                });

                if (matchingParagraphsCount === 0 && (curFilterMode === 'failed' || curSearchQuery)) {
                    // Skip showing this file if no matching rows
                    return;
                }

                renderedFilesCount++;

                const fileBlock = document.createElement('div');
                fileBlock.className = 'bg-white rounded-2xl border border-slate-200 overflow-hidden shadow-sm animate-fade-in';
                
                fileBlock.innerHTML = `
                    <div class="px-6 py-4.5 bg-slate-50/70 border-b border-slate-100 flex flex-col sm:flex-row sm:items-center justify-between gap-3">
                        <div class="flex items-center gap-3">
                            <div class="p-2 bg-indigo-50 rounded-lg text-indigo-600">
                                <svg class="w-5 h-5" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M7 21h10a2 2 0 002-2V9.414a1 1 0 00-.293-.707l-5.414-5.414A1 1 0 0012.586 3H7a2 2 0 00-2 2v14a2 2 0 002 2z" /></svg>
                            </div>
                            <div>
                                <h3 class="text-sm font-bold text-slate-800">${f.fileName}</h3>
                                <p class="text-[10.5px] text-slate-400 font-medium font-mono mt-0.5">Size: ${(f.fileSize/1024).toFixed(1)} KB | Path: ${escapeHtml(f.filePath)}</p>
                            </div>
                        </div>
                        <div class="flex items-center gap-1.5 self-start sm:self-center">
                            <span class="inline-flex items-center px-2.5 py-0.5 rounded-full text-[10px] font-bold bg-indigo-50 border border-indigo-100 text-indigo-600">
                                Checked Paragraphs: ${f.summary.totalParagraphs}
                            </span>
                            <span class="inline-flex items-center px-2.5 py-0.5 rounded-full text-[10px] font-bold bg-indigo-50 border border-indigo-100 ${f.summary.failed > 0 ? 'bg-rose-50 text-rose-600 border-rose-100 font-bold' : 'bg-emerald-50 text-emerald-600 border-emerald-100'}">
                                Validation Faults: ${f.summary.failed}
                            </span>
                        </div>
                    </div>
                    <div class="overflow-x-auto">
                        <table class="w-full text-left border-collapse">
                            <thead>
                                <tr class="bg-slate-50/30 border-b border-slate-100 text-slate-400 text-[10px] font-bold tracking-wider font-sans uppercase">
                                    <th class="px-5 py-3 w-16 text-center">Num</th>
                                    <th class="px-5 py-3">Content Snippet</th>
                                    <th class="px-5 py-3 w-48">Extracted Format</th>
                                    <th class="px-5 py-3 w-36">Expected Level</th>
                                    <th class="px-5 py-3 w-24 text-center">Status</th>
                                    <th class="px-5 py-3 w-64">Validation Deviations</th>
                                </tr>
                            </thead>
                            <tbody class="divide-y divide-slate-50 font-sans">
                                ${fileRowsHtml}
                            </tbody>
                        </table>
                    </div>
                `;
                container.appendChild(fileBlock);
            });

            if (renderedFilesCount === 0) {
                container.innerHTML = `
                    <div class="p-12 text-center bg-white rounded-2xl border border-slate-150 shadow-xs">
                        <svg class="w-8 h-8 text-slate-300 mx-auto mb-2" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9.172 16.172a4 4 0 015.656 0M9 10h.01M15 10h.01M21 12a9 9 0 11-18 0 9 9 0 0118 0z" /></svg>
                        <p class="text-xs text-slate-400 font-medium">No files or segments matching current filters.</p>
                    </div>
                `;
            }
        }

        function escapeHtml(str) {
            if (!str) return '';
            return str.replace(/&/g, "&amp;").replace(/</g, "&lt;").replace(/>/g, "&gt;").replace(/"/g, "&quot;").replace(/'/g, "&#039;");
        }
    </script>
</body>
</html>"""

@app.route("/")
def home():
    return render_template_string(HTML_TEMPLATE)

@app.route("/api/analyze", methods=["POST"])
def run_analysis():
    data = request.get_json() or {}
    folder_path = data.get("folder_path", "").strip()
    output_path = data.get("output_path", "").strip() or None
    
    if not folder_path:
        return jsonify({"error": "Folder path is required."}), 400
        
    reports, err = scan_and_analyze_folder(folder_path, output_path)
    if err:
        return jsonify({"error": err}), 400
        
    return jsonify({
        "success": True,
        "reports": reports
    })

if __name__ == "__main__":
    print("-" * 55)
    print("DOCX Standards Analyzer local Python server starter...")
    print("To run correctly, please type 'pip install flask python-docx lxml' first.")
    print("-" * 55)
    app.run(host="127.0.0.1", port=5000, debug=True)
